import os
import json
import logging
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from docx import Document
from io import BytesIO
import streamlit as st

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

class DocumentProcessor:
    """
    Enhanced document processor supporting PDF, images, and Word documents.
    Uses OCR for text extraction with improved preprocessing.
    """
    
    def __init__(self):
        self.supported_formats = {
            'pdf': ['.pdf'],
            'image': ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif'],
            'word': ['.docx', '.doc']
        }
    
    def get_file_type(self, filename):
        """Determine file type based on extension."""
        ext = os.path.splitext(filename.lower())[1]
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        return None
    
    def preprocess_image(self, image):
        """
        Preprocess image for better OCR results.
        """
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast and sharpness
            from PIL import ImageEnhance, ImageFilter
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            # Sharpen image
            image = image.filter(ImageFilter.SHARPEN)
            
            return image
        except Exception as e:
            logging.warning(f"Image preprocessing failed: {e}")
            return image
    
    def extract_text_from_image(self, image, page_number=1):
        """
        Extract text from image using OCR with preprocessing.
        """
        try:
            # Preprocess image
            processed_image = self.preprocess_image(image)
            
            # OCR configuration for better accuracy
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz!@#$%^&*()_+-=[]{}|;:,.<>?/~` '
            
            # Extract text
            text = pytesseract.image_to_string(processed_image, config=custom_config).strip()
            
            if not text:
                logging.warning(f"No text extracted from page {page_number}")
                return ""
            
            return text
            
        except Exception as e:
            logging.error(f"OCR extraction failed for page {page_number}: {e}")
            return ""
    
    def process_pdf(self, file_content, filename):
        """
        Process PDF file: convert to images and extract text.
        """
        try:
            with st.status("Processing PDF...", expanded=True) as status:
                st.write("Converting PDF to images...")
                
                # Convert PDF to images
                images = convert_from_path(BytesIO(file_content), dpi=150)
                
                st.write(f"Processing {len(images)} pages...")
                
                metadata = []
                progress_bar = st.progress(0)
                
                for i, image in enumerate(images):
                    page_number = i + 1
                    
                    # Update progress
                    progress = (i + 1) / len(images)
                    progress_bar.progress(progress)
                    st.write(f"Extracting text from page {page_number}...")
                    
                    # Extract text using OCR
                    text = self.extract_text_from_image(image, page_number)
                    
                    # Store metadata
                    metadata.append({
                        "page": page_number,
                        "text": text,
                        "source_file": filename,
                        "file_type": "pdf"
                    })
                
                status.update(label="PDF processing completed!", state="complete")
                
            return metadata
            
        except Exception as e:
            st.error(f"Error processing PDF: {e}")
            logging.error(f"PDF processing error: {e}")
            return []
    
    def process_image(self, file_content, filename):
        """
        Process single image file.
        """
        try:
            with st.status("Processing Image...", expanded=True) as status:
                st.write("Loading image...")
                
                # Load image
                image = Image.open(BytesIO(file_content))
                
                st.write("Extracting text from image...")
                
                # Extract text
                text = self.extract_text_from_image(image)
                
                metadata = [{
                    "page": 1,
                    "text": text,
                    "source_file": filename,
                    "file_type": "image"
                }]
                
                status.update(label="Image processing completed!", state="complete")
                
            return metadata
            
        except Exception as e:
            st.error(f"Error processing image: {e}")
            logging.error(f"Image processing error: {e}")
            return []
    
    def process_word(self, file_content, filename):
        """
        Process Word document.
        """
        try:
            with st.status("Processing Word Document...", expanded=True) as status:
                st.write("Extracting text from Word document...")
                
                # Load Word document
                doc = Document(BytesIO(file_content))
                
                # Extract text from paragraphs
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                text = "\n".join(full_text)
                
                metadata = [{
                    "page": 1,
                    "text": text,
                    "source_file": filename,
                    "file_type": "word"
                }]
                
                status.update(label="Word document processing completed!", state="complete")
                
            return metadata
            
        except Exception as e:
            st.error(f"Error processing Word document: {e}")
            logging.error(f"Word processing error: {e}")
            return []
    
    def process_document(self, uploaded_file):
        """
        Main method to process any supported document type.
        """
        try:
            filename = uploaded_file.name
            file_content = uploaded_file.read()
            file_type = self.get_file_type(filename)
            
            if not file_type:
                st.error(f"Unsupported file format: {filename}")
                return []
            
            st.info(f"Processing {file_type.upper()} file: {filename}")
            
            if file_type == 'pdf':
                return self.process_pdf(file_content, filename)
            elif file_type == 'image':
                return self.process_image(file_content, filename)
            elif file_type == 'word':
                return self.process_word(file_content, filename)
            else:
                st.error(f"Processing not implemented for {file_type}")
                return []
                
        except Exception as e:
            st.error(f"Error processing document: {e}")
            logging.error(f"Document processing error: {e}")
            return []
    
    def save_metadata(self, metadata, metadata_file):
        """
        Save processed metadata to JSON file.
        """
        try:
            os.makedirs(os.path.dirname(metadata_file), exist_ok=True)
            
            with open(metadata_file, "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=4, ensure_ascii=False)
            
            logging.info(f"Metadata saved to {metadata_file}")
            return True
            
        except Exception as e:
            logging.error(f"Error saving metadata: {e}")
            return False